from datetime import datetime
from pathlib import Path
from docx import Document
from anonymator.output_naming import anonymized_path
from anonymator.report.audit import AuditReport
from anonymator.files.ooxml import scan, xml_parts
from anonymator.files.ooxml.text_unit import TextUnit
from anonymator.files.ooxml.xml_parts import XmlRun

_W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def _iter_block(container, prefix: str):
    for p in container.paragraphs:
        if p.runs:
            yield TextUnit(list(p.runs), prefix)
    for table in container.tables:
        yield from _iter_table(table, prefix)


def _iter_table(table, prefix: str):
    base = "" if prefix == "Corps" else f"{prefix} / "
    for ri, row in enumerate(table.rows, 1):
        for ci, cell in enumerate(row.cells, 1):
            loc = f"{base}Tableau L{ri}C{ci}"
            for p in cell.paragraphs:
                if p.runs:
                    yield TextUnit(list(p.runs), loc)
            for nested in cell.tables:
                yield from _iter_table(nested, loc)


def _iter_textboxes(doc):
    t_tag = f"{{{_W}}}t"
    body = doc.element.body
    for txbx in body.iter(f"{{{_W}}}txbxContent"):
        for p in txbx.iter(f"{{{_W}}}p"):
            runs = [XmlRun(r, t_tag) for r in p.findall(f"{{{_W}}}r")]
            if runs:
                yield TextUnit(runs, "Zone de texte")


def iter_main_units(doc):
    """Unités des conteneurs de la partie principale (sauvegardées nativement
    par doc.save) : corps, tableaux, en-têtes/pieds, zones de texte."""
    yield from _iter_block(doc, "Corps")
    for section in doc.sections:
        if not section.header.is_linked_to_previous:
            yield from _iter_block(section.header, "En-tête")
        if not section.footer.is_linked_to_previous:
            yield from _iter_block(section.footer, "Pied")
    yield from _iter_textboxes(doc)


def anonymize_document(path: Path, ner, ref, output_dir: Path,
                       when: datetime) -> tuple[Path, AuditReport]:
    doc = Document(str(path))
    units = list(iter_main_units(doc))
    retained = scan.confirmed_only(scan.scan_units(units, ner, ref))
    report = scan.apply_units(units, retained, ref)
    out = anonymized_path(path, output_dir, when)
    doc.save(str(out))
    xml_parts.postprocess_docx(out, ner, ref, report)
    return out, report
