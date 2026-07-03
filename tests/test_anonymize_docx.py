from datetime import datetime
import zipfile
from docx import Document
from anonymator.referential import Referential
from anonymator.ner import FakeNer
from anonymator.files.ooxml import docx_io
from tests.ooxml_fixtures import make_docx


def _anonymize(tmp_path):
    src = make_docx(tmp_path / "src.docx")
    ref = Referential.load_default()
    ner = FakeNer({"Claire Martin": "PERSON"})
    out, report = docx_io.anonymize_document(
        src, ner, ref, tmp_path, when=datetime(2026, 7, 3, 9, 0, 0))
    return src, out, report


def test_output_name_and_body_masked(tmp_path):
    _, out, _ = _anonymize(tmp_path)
    assert out.name == "src_ano_20260703090000.docx"
    doc = Document(str(out))
    assert doc.paragraphs[0].text == "Contact : [PERSONNE]"


def test_table_header_footer_masked(tmp_path):
    _, out, _ = _anonymize(tmp_path)
    doc = Document(str(out))
    assert doc.tables[0].rows[0].cells[1].text == "[PERSONNE]"
    assert "[PERSONNE]" in doc.sections[0].header.paragraphs[0].text
    assert "[PERSONNE]" in doc.sections[0].footer.paragraphs[0].text


def test_textbox_masked(tmp_path):
    # La zone de texte vit dans document.xml : elle survit au round-trip
    # python-docx. (Commentaires/notes = testés dans test_ooxml_xml_parts.py.)
    _, out, _ = _anonymize(tmp_path)
    with zipfile.ZipFile(out) as z:
        document = z.read("word/document.xml").decode("utf-8")
    assert "Claire Martin" not in document and "[PERSONNE]" in document


def test_metadata_purged(tmp_path):
    _, out, _ = _anonymize(tmp_path)
    doc = Document(str(out))
    assert (doc.core_properties.author or "") == ""
    assert (doc.core_properties.last_modified_by or "") == ""


def test_report_and_original_untouched(tmp_path):
    src, _, report = _anonymize(tmp_path)
    assert any(r["original"] == "Claire Martin" for r in report.to_rows())
    assert "Claire Martin" in Document(str(src)).paragraphs[0].text


def test_nested_table_masked(tmp_path):
    # Tableau imbriqué dans une cellule (récursion de _iter_table).
    doc = Document()
    outer = doc.add_table(rows=1, cols=1)
    inner = outer.rows[0].cells[0].add_table(rows=1, cols=1)
    inner.rows[0].cells[0].text = "Client Claire Martin"
    src = tmp_path / "nested.docx"
    doc.save(str(src))
    out, report = docx_io.anonymize_document(
        src, FakeNer({"Claire Martin": "PERSON"}),
        Referential.load_default(), tmp_path, when=datetime(2026, 7, 3, 9, 0, 0))
    d = Document(str(out))
    inner_cell = d.tables[0].rows[0].cells[0].tables[0].rows[0].cells[0]
    assert inner_cell.text == "Client [PERSONNE]"
    # La localisation reflète l'imbrication (préfixe du tableau parent).
    locs = [r["locations"] for r in report.to_rows()]
    assert any("Tableau L1C1 / Tableau L1C1" in loc for loc in locs)
